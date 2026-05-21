import os
import datetime
import numpy as np
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import warnings

def init_document_style(document):
    style = document.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor(0, 0, 0)

def init_page_style(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

def init_header_style(document):
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run("银联智策顾问（上海）有限公司")
    run.bold = True
    run.font.size = Pt(10.5)
    paragraph.add_run().add_picture(
        logo_path,
        width=Inches(0.2),
        height=Inches(0.14)
    )
    paragraph.alignment = 1

def init_footer_style(document):

    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run("注：统计结果仅供参考")
    run.bold = True
    run.font.size = Pt(10)
    paragraph.alignment = 1
