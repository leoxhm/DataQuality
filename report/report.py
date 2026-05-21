import os
import datetime
import numpy as np
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import warnings

warnings.filterwarnings('ignore')
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False

data_processor_dict = {"liyue.bai":"白礼月", "zedong.cai":"蔡泽栋", "zijian.wang":"王子剑",\
                       "lu.cao":"曹璐","zhixiang.chen":"陈志祥","jiaying.fang":"方佳盈",\
                       "xinlei.ge":"葛心蕾","jingjun.hu":"胡静珺","ming.le":"乐明",\
                       "wei.li":"李维","yebei.li":"李业北","haitao.liu":"刘海涛",\
                       "xiaohan.mei":"梅笑寒","chen.meng":"孟晨","yiming.niu":"牛怡鸣",\
                       "ziqi.qin":"秦子淇","xiaoduo.sang":"桑小朵","qian.shen":"沈倩",\
                       "ziqi.sheng":"盛子奇","jiawei.sun":"孙家伟","jiamao.sun":"孙嘉懋",\
                       "dacheng.wang":"王大成","haotian.wei":"魏昊天","xuefei.wu":"吴雪菲",\
                       "tingting.xu":"徐婷婷","wang.yu":"虞王","haixia.yuan":"袁海霞",\
                       "letao.zhang":"张乐涛","yaping.zhang":"张亚平","zewen.zhang":"张则文",\
                       "zhuangfei.zhao":"赵壮飞","ruichen.zhou":"周睿辰",\
                       "yunhao.zhu":"朱云皓","youwen.zuo":"左又文"
                    }

product_eng_to_chn = {"wealth_score_res":"综合财富分", "multi_key_vars":"衍生变量v3","hnsk_res":"海纳定制",\
                      "hirisk":"周度风险","hirisk_ext":"周度衍生","cash_out":"新版非正常消费",\
                      "hknl_basic":"还款能力基础版","dz_monitor":"贷中监控","js_monitor":"金哨监控",\
                      "hknl_monitor":"还款能力监控版","tq_bank_score":"天权分银行版","tq_bank1_score":"天权分银行版v1.1",\
                      "tq_xj_score":"天权分消金版","tq_lhd_xfd_score":"天权分联合贷消费贷版"
                      }

gp_col_dict = {"综合财富分":"yearmonth", "衍生变量v3":"yearmonth","海纳定制":"yearmonth",\
              "周度风险":"week_dt","周度衍生":"week_dt","新版非正常消费":"yearmonth",\
              "还款能力基础版":"week_dt","贷中监控":"yearmonnth","金哨监控":"yearmonth",\
              "还款能力监控版":"yearmonth","天权分银行版":"yearmonth","天权分银行版v1.1":"yearmonth",\
              "天权分消金版":"yearmonth","天权分联合贷消费贷版":"yearmonth"
              }

document = Document()
document.styles["Normal"].font.name = "微软雅黑"
document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
document.styles["Normal"].font.size = Pt(10.5)
document.styles["Normal"].font.color.rgb = RGBColor(0,0,0)

section = document.sections[0]
header = section.header
header_paragraph = header.paragraphs[0]
header_paragraph.add_run("银联智策顾问（上海）有限公司").bold=True
header_paragraph.font.size = Pt(10.5)
header_paragraph.add_run().add_picture(log_path, width=Inches(0.2), height=Inches(0.14))
header_paragraph.style.font.size = Pt(10)
header.paragraph.alignment = 1

footer = section.footer
footer_paragraph = footer.paragraphs[0]
footer_paragraph.add_run("z").bold = True
footer_paragraph.font.size = Pt(10)
footer_paragraph.alignment = 1

section = document.sections[0]
section.page_width = docx.shared.Inches(8.5)
section.page_height  =docx.shared.Inches(11)
section.top_margin = docx.shared.Inches(1)
section.bottom_margin = docx.shared.Inches(1)

document.add_heading("数据QC简报", 0)
p = document.add_paragraph('提数结果生产于')