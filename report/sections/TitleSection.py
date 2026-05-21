import os
import datetime
import numpy as np
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import warnings
from ..config import data_processor_dict

class TitleSection:
    def render(self, document, ctx):
        document.add_heading("数据QC简报", 0)
        p = document.add_paragraph('提数结果生产于')
        p.add_run(datetime.datetime.today().strftime('%Y-%m-%d')).bold = True
        p.add_run('    提数人员：')
        p.add_run(data_processor_dict[ctx.config.name_domain]).italic = True
