import os
import datetime
import numpy as np
import pandas as pd
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import warnings

from pptx.oxml.xmlchemy import OxmlElement

from ..config import data_processor_dict, product_eng_to_chn, color_dict, monitor_product

class AbstractSection:
    def render(self, document, ctx):
        document.add_heading("产品质量综述", 1)
        document.add_picture()
        document.add_paragraph("该部分综述产品", style = "Intense Quote")

        p = document.add_paragraph(style='List Bullet')
        run = p.add_run("待确认事项：")
        run.bold = True
        if product_eng_to_chn[ctx.config.table_nm] not in monitor_product:
            p.add_run(" 核对三体单号、筛选字段、拼接其他产品")
        else:
            p.add_run(" 核对三体单号")

        p = document.add_paragraph(style='List Bullet')
        run = p.add_run(f"是否有时效：{ctx.metric[""]}")
        run.bold = True
        run.font.color.rgb = color_dict[]
        p.add_run(f" 样本上传与样本上车间隔{}天，样本上车与样本下车间隔{}天")

        p = document.add_paragraph(style='List Bullet')
        run = p.add_run(f"是否异常：{ctx.metric[""]}")
        run.bold = True
        run.font.color.rgb = color_dict[ctx.metric[""]]
        p.add_run(f" 无差异字段数{}(详细请见报告变量探测部分)")

        p = document.add_paragraph(style='List Bullet')
        run = p.add_run(f"是否完整：{ctx.metric[""]}")
        run.bold = True
        run.font.color.rgb = color_dict[ctx.metric[""]]
        if product_eng_to_chn[ctx.config.table_nm] not in monitor_product:
            complete_table(document, [ctx.metric[""], ctx.metric[""], ctx.metric[""], ctx.metric[""]])
        else:
            pass

        p = document.add_paragraph(style='List Bullet')
        run = p.add_run(f"是否稳定：{ctx.metric[""]}")
        run.bold = True
        run.font.color.rgb = color_dict[ctx.metric[""]]
        stable_table(document, [ctx.metric[""], ctx.metric[""], ctx.metric[""]])

    def complete_table(self, doc, rows_data):
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Medium Shading 1 Accent 1'
        headers = ['下车字段数','','','']
        rows_data = [rows_data]

        for j, header in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = header

        for i, row_data in enumerate(rows_data, start=1):
            for j, value in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = ""
                p1 = cell.paragraphs[0]
                if j >= 1:
                    run1 = p1.add_run(f"{value:.2%}")
                elif j ==0:
                    run1 = p1.add_run(str(value))
                else:
                    run1 = p1.add_run("")
                run1.font.size = Pt(10)
                run1.bold = True
                p1.paragraph_format.space_after = Pt(1)

                p2 = cell.add_paragraph()
                if j==1:
                    run2 = p2.add_run("(参考范围：【0%， 50%））")
                elif j==2:
                    run2 = p2.add_run("(参考范围：【0%， 25%））")
                elif j==3:
                    run2 = p2.add_run("(参考范围：【0%， 10%））")
                elif j==0:
                    run2 = p2.add_run("(请核对需求字典）")
                else:
                    run2 = p2.add_run("")
                run1.font.size = Pt(8)
                run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                p2.paragraph_format.space_before = Pt(0)

                try:
                    if j == 1 and float(value) > 0.5:
                        add_arrow(table_cell(i, j), '', RGBColor(0xFF,0x00,0x00))
                    elif j == 2 and float(value) > 0.25:
                        add_arrow(table_cell(i, j), '', RGBColor(0xFF,0x00,0x00))
                    elif j == 3 and float(value) > 0.10:
                        add_arrow(table_cell(i, j), '', RGBColor(0xFF,0x00,0x00))
                    else:
                        continue
                except:
                    pass

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:vAlign_r'), str(cell.alignment))
                    tcPr.append(vAlign)
                    for run in paragraph.runs:
                        run.font.name = "微软雅黑"
                        run.font.size = Pt(10)

        for j in range(len(headers)):
            for paragraph in table.cell(0, j).paragraphs:
                for run in paragraph.runs:
                    run.font.bold =True

        remarks = [
            "口径1: P(字段缺失率50%) = 缺失率不低于50%字段数/下车字段数",
            "口径2: P(字段缺失率80%) = 缺失率不低于80%字段数/下车字段数",
            "口径3: P(字段缺失率99.9%) = 缺失率不低于99.9%字段数/下车字段数"
        ]
        add_remark_row(table, remarks, remark_label = "备注*")


    def add_arrow(self, cell, arrow_char, color):
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            run = paragraph.add_run(arrow_char)
            if color:
                run.font.color.rgb = color

    def add_remark_row(self, table, remark_texts, remark_label):
        new_row = table.add_row()

        first_cell = new_row.cells[0]
        first_cell.text = remark_label
        for paragraph in first_cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

        if len(new_row.cells) > 2:
            start_cell = new_row.cells[1]
            end_cell = new_row.cells[-1]
            start_cell.merge(end_cell)
            merged_cell = start_cell
        else:
            merged_cell = new_row.cells

        merged_cell.text = ""
        for idx,text in enumerate(remark_texts):
            if idx == 0:
                p = merged_cell.paragraphs[0]
            else:
                p = merged_cell.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
